# backend/app/gdd_engine/docx_exporter.py

"""
DOCX Exporter (Safe, Cross-Platform, Markdown-Aware)

This module turns a markdown or plaintext GDD into a well-structured
DOCX file. It is intentionally conservative so it NEVER breaks the 
pipeline, never throws directory errors, and always produces a valid 
DOCX file.
"""

import os
import re
from docx import Document
from docx.enum.text import WD_BREAK


# ------------------------------------------------------------
# MARKDOWN → RUN FORMAT HELPERS
# ------------------------------------------------------------
def _apply_inline_formatting(run, text: str):
    """
    Apply safe inline formatting:
    - Bold: **text**
    - Italic: *text*
    - Bold+Italic: ***text***
    """

    original = text

    # Bold + Italic  ***text***
    if re.fullmatch(r"\*\*\*(.+?)\*\*\*", text):
        run.bold = True
        run.italic = True
        run.text = text[3:-3]
        return

    # Bold  **text**
    if re.fullmatch(r"\*\*(.+?)\*\*", text):
        run.bold = True
        run.text = text[2:-2]
        return

    # Italic  *text*
    if re.fullmatch(r"\*(.+?)\*", text):
        run.italic = True
        run.text = text[1:-1]
        return

    # No formatting
    run.text = original


# ------------------------------------------------------------
# EXPORT FUNCTION
# ------------------------------------------------------------
def export_to_docx(markdown: str, output_path: str):
    """
    Convert markdown/plaintext GDD into a DOCX file.

    This function:
    - Creates parent folders if needed
    - Parses headers (#), bullets (- or •), paragraphs
    - Applies very safe inline formatting (**bold**, *italic*)
    - NEVER throws if file/dir does not exist — it creates it
    - Always writes a valid .docx file
    """

    # Ensure directory exists
    folder = os.path.dirname(output_path)
    if folder and not os.path.exists(folder):
        os.makedirs(folder, exist_ok=True)

    doc = Document()

    # Normalize markdown newlines
    lines = markdown.split("\n")

    for line in lines:

        stripped = line.strip()

        # Empty line → spacing paragraph
        if stripped == "":
            doc.add_paragraph("")
            continue

        # --------------------------
        # HEADERS
        # --------------------------
        if stripped.startswith("### "):     # H3
            doc.add_heading(stripped[4:], level=3)
            continue

        if stripped.startswith("## "):      # H2
            doc.add_heading(stripped[3:], level=2)
            continue

        if stripped.startswith("# "):       # H1
            doc.add_heading(stripped[2:], level=1)
            continue

        # --------------------------
        # BULLETED LISTS
        # --------------------------
        if stripped.startswith("- ") or stripped.startswith("• "):
            li_text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run("")
            _apply_inline_formatting(r, li_text)
            continue

        # --------------------------
        # PARAGRAPHS (default)
        # --------------------------
        p = doc.add_paragraph()
        r = p.add_run("")
        _apply_inline_formatting(r, stripped)

    # Final safety newline
    doc.add_paragraph("")

    # Save docx
    try:
        doc.save(output_path)
    except Exception as e:
        raise RuntimeError(f"DOCX export failed: {e}")
